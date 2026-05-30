# -*- coding: utf-8 -*-
"""
Generateur du Cahier des Charges (CDC) - Systeme de Gestion de Projets et de Taches
Produit un document .docx professionnel et structure.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Palette de l'application (charte graphique entreprise)
# ---------------------------------------------------------------------------
NAVY      = RGBColor(0x1A, 0x27, 0x44)   # Bleu marine profond
BLUE      = RGBColor(0x25, 0x63, 0xEB)   # Bleu corporate
LIGHT_BG  = "EAF0F8"                       # Fond clair (hex string pour shading)
NAVY_HEX  = "1A2744"
BLUE_HEX  = "2563EB"
GREY      = RGBColor(0x4B, 0x55, 0x63)
GREEN     = RGBColor(0x10, 0xB9, 0x81)
ORANGE    = RGBColor(0xF5, 0x9E, 0x0B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Calibri"

doc = Document()

# ---------------------------------------------------------------------------
# Styles de base
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x21, 0x29, 0x33)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_bottom_border(p, color_hex=BLUE_HEX, size=18):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def h1(text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run((f"{number}.  " if number else "") + text)
    run.font.name = FONT
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = NAVY
    add_bottom_border(p, BLUE_HEX, 18)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE
    return p


def para(text, bold=False, italic=False, color=None, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.name = FONT
        r.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.name = FONT
    else:
        r = p.add_run(text)
        r.font.name = FONT
    return p


def make_table(headers, rows, col_widths=None, header_bg=NAVY_HEX, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_bg(hdr[i], header_bg)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(htext)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT
        run.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            set_cell_margins(cells[ci])
            if zebra and ri % 2 == 0:
                set_cell_bg(cells[ci], LIGHT_BG)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            # support tuple (text, color) for first-col badges
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = FONT
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ===========================================================================
# PAGE DE GARDE
# ===========================================================================
# Bandeau haut
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.3)
sec.right_margin = Cm(2.3)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CAHIER DES CHARGES")
r.font.size = Pt(34)
r.font.bold = True
r.font.color.rgb = NAVY
r.font.name = FONT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, BLUE_HEX, 28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
r = p.add_run("Plateforme de Gestion de Projets, de Taches\net de Collaboration d'Equipe")
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = BLUE
r.font.name = FONT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Application web full-stack — Spring Boot & Angular 21")
r.font.size = Pt(13)
r.font.italic = True
r.font.color.rgb = GREY
r.font.name = FONT

for _ in range(4):
    doc.add_paragraph()

# Bloc d'identification
info = doc.add_table(rows=0, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info_rows = [
    ("Intitule du projet", "Systeme de Gestion de Projets et de Taches (GPI App)"),
    ("Type de document", "Cahier des charges fonctionnel et technique"),
    ("Version", "1.0"),
    ("Date", "29 mai 2026"),
    ("Statut du produit", "Operationnel — en evolution continue"),
    ("Etablissement", "Saint Jean — Ecole d'Ingenieurs"),
    ("Auteur", "Equipe de developpement"),
]
for k, v in info_rows:
    cells = info.add_row().cells
    set_cell_bg(cells[0], NAVY_HEX)
    set_cell_margins(cells[0]); set_cell_margins(cells[1])
    set_cell_bg(cells[1], LIGHT_BG)
    rp = cells[0].paragraphs[0].add_run(k)
    rp.font.bold = True; rp.font.color.rgb = WHITE; rp.font.size = Pt(10.5); rp.font.name = FONT
    rv = cells[1].paragraphs[0].add_run(v)
    rv.font.size = Pt(10.5); rv.font.name = FONT
    cells[0].width = Inches(2.2); cells[1].width = Inches(4.3)

doc.add_page_break()

# ===========================================================================
# SOMMAIRE (statique)
# ===========================================================================
h1("Sommaire")
toc_items = [
    "1.  Presentation generale du projet",
    "2.  Contexte et problematique",
    "3.  Objectifs du systeme",
    "4.  Perimetre fonctionnel et acteurs",
    "5.  Etat des lieux — Fonctionnalites realisees",
    "6.  Architecture technique existante",
    "7.  Axes d'innovation a developper",
    "8.  Fonctionnalites a ajouter (backlog)",
    "9.  Exigences non fonctionnelles",
    "10. Planning previsionnel et lots de livraison",
    "11. Risques et mesures de mitigation",
    "12. Criteres de reception et conclusion",
]
for t in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(t)
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY
    r.font.name = FONT

doc.add_page_break()

# ===========================================================================
# 1. PRESENTATION GENERALE
# ===========================================================================
h1("Presentation generale du projet", "1")
para("Le present document constitue le cahier des charges de la plateforme de gestion de "
     "projets et de taches developpee dans le cadre du systeme GPI App. Il decrit le perimetre "
     "fonctionnel, l'architecture technique, l'etat d'avancement du produit ainsi que les axes "
     "d'innovation et d'evolution retenus.")
para("Il s'agit d'une application web full-stack destinee a centraliser le pilotage des projets, "
     "l'organisation des taches, la collaboration des equipes et le suivi de la performance au sein "
     "d'une organisation. Le systeme repose sur une architecture moderne decouplee : une API REST "
     "Spring Boot cote serveur et une application monopage (SPA) Angular 21 cote client.")

h2("Vision produit")
para("Offrir un environnement de travail unifie, professionnel et evolutif, permettant aux "
     "administrateurs, chefs de projet et collaborateurs de planifier, executer et suivre l'ensemble "
     "des activites d'un projet depuis une interface unique, securisee et en temps reel.")

# ===========================================================================
# 2. CONTEXTE ET PROBLEMATIQUE
# ===========================================================================
h1("Contexte et problematique", "2")
para("La gestion des projets dans de nombreuses organisations souffre encore d'une dispersion des "
     "outils : suivi des taches dans un tableur, echanges par messagerie externe, depot des livrables "
     "par courriel, reporting manuel. Cette fragmentation entraine une perte d'information, un manque "
     "de visibilite sur l'avancement reel et une difficulte a mesurer la performance des equipes.")
para("Le projet GPI App repond a ce besoin en proposant une solution integree couvrant l'ensemble du "
     "cycle de vie d'un projet :")
for t in [
    "centralisation des projets, taches, equipes et livrables ;",
    "collaboration en temps reel (messagerie interne, notifications instantanees) ;",
    "tracabilite complete des actions (journal d'activite) ;",
    "pilotage par indicateurs (tableaux de bord et rapports exportables).",
]:
    bullet(t)

# ===========================================================================
# 3. OBJECTIFS
# ===========================================================================
h1("Objectifs du systeme", "3")
make_table(
    ["Objectif", "Description"],
    [
        ["Centralisation", "Regrouper projets, taches, equipes, livrables et echanges dans une plateforme unique."],
        ["Collaboration", "Faciliter le travail d'equipe via messagerie, commentaires et notifications temps reel."],
        ["Pilotage", "Fournir des tableaux de bord et rapports pour suivre l'avancement et la performance."],
        ["Securite", "Garantir un acces controle par roles avec authentification forte (JWT, BCrypt)."],
        ["Tracabilite", "Journaliser les actions sensibles et les tentatives de connexion."],
        ["Evolutivite", "Reposer sur une architecture decouplee et modulaire facilitant l'ajout de modules."],
    ],
    col_widths=[1.7, 4.8],
)

# ===========================================================================
# 4. PERIMETRE FONCTIONNEL ET ACTEURS
# ===========================================================================
h1("Perimetre fonctionnel et acteurs", "4")
para("Le systeme repose sur un modele de controle d'acces base sur les roles (RBAC). Trois profils "
     "d'utilisateurs disposent chacun d'un espace de travail (layout) et de droits distincts.")
make_table(
    ["Acteur", "Role et droits principaux"],
    [
        ["Administrateur",
         "Acces total : gestion des utilisateurs, des projets, configuration systeme, supervision de "
         "l'activite, securite (tentatives de connexion), rapports globaux et support."],
        ["Chef de projet (PM)",
         "Creation et pilotage des projets, constitution des equipes, delegation et suivi des taches, "
         "revue des livrables, calendrier, rapports de projet."],
        ["Collaborateur (Utilisateur)",
         "Consultation et mise a jour de ses taches, depot des livrables, saisie des temps, messagerie "
         "interne, calendrier personnel et tickets de support."],
    ],
    col_widths=[1.9, 4.6],
)

# ===========================================================================
# 5. ETAT DES LIEUX - REALISE
# ===========================================================================
h1("Etat des lieux — Fonctionnalites realisees", "5")
para("Cette section recense les modules deja developpes et operationnels au sein de la plateforme. "
     "Le produit est actuellement fonctionnel et couvre l'ensemble du socle de gestion de projet.",
     italic=True, color=GREY)

modules = [
    ("Authentification & securite",
     "Inscription, connexion par JWT, hachage BCrypt, reinitialisation de mot de passe par email, "
     "suivi des tentatives de connexion, gestion des sessions et controle d'acces par roles."),
    ("Gestion des utilisateurs",
     "CRUD complet des comptes, affectation des roles, activation/desactivation, profils."),
    ("Gestion des projets",
     "Creation, modification, suivi d'avancement, page de detail enrichie (vue d'ensemble, graphiques, "
     "sante du projet, jalons, allocation des equipes)."),
    ("Gestion des taches",
     "Creation, affectation, priorites, echeances, suivi en pourcentage, tableau Kanban, commentaires "
     "et filtres."),
    ("Gestion des equipes",
     "Constitution d'equipes, affectation des membres aux projets, repartition de la charge."),
    ("Livrables",
     "Depot de fichiers, workflow de revue et d'approbation (en attente / approuve / rejete), retours."),
    ("Suivi du temps",
     "Saisie des heures par tache, recapitulatifs (jour / semaine / mois)."),
    ("Calendrier",
     "Vue type Google Agenda (mois / semaine / jour / liste), evenements et echeances, configuration "
     "Google Calendar prevue cote serveur."),
    ("Messagerie interne",
     "Discussions individuelles en temps reel (interface type WhatsApp Web), indicateurs de messages "
     "non lus."),
    ("Notifications temps reel",
     "Diffusion via WebSocket, cloche de notification, marquage lu/non-lu, preferences de notification "
     "par utilisateur."),
    ("Tableaux de bord",
     "Indicateurs cles et graphiques interactifs (Chart.js) distincts par role."),
    ("Rapports & exports",
     "Generation de rapports (modeles JasperReports / JRXML) et export de donnees."),
    ("Support",
     "Systeme de tickets de support integre (creation, suivi, traitement)."),
    ("Journal d'activite",
     "Tracabilite des actions sensibles pour audit et supervision."),
    ("Documentation API",
     "Documentation interactive Swagger / OpenAPI 3.0 de l'ensemble des endpoints REST."),
    ("Charte graphique entreprise",
     "Refonte UI complete : palette 3 couleurs (marine / bleu / clair), 35 feuilles de style, "
     "interface responsive et homogene."),
]
make_table(
    ["Module realise", "Description"],
    [[m[0], m[1]] for m in modules],
    col_widths=[1.9, 4.6],
)

para("Synthese : le socle metier est complet. Les 16 entites principales (Utilisateurs, Projets, "
     "Taches, Equipes, Livrables, Messages, Notifications, Evenements, Commentaires, Temps, Tickets, "
     "Journaux, etc.) sont implementees et exposees via une API REST documentee.",
     bold=True, color=NAVY)

# ===========================================================================
# 6. ARCHITECTURE TECHNIQUE
# ===========================================================================
h1("Architecture technique existante", "6")
make_table(
    ["Couche", "Technologies"],
    [
        ["Backend (API)", "Spring Boot 3.x, Spring Security, Spring Data JPA, architecture en couches "
                           "(controller / service / repository / DTO / entity)."],
        ["Securite", "JWT (jetons), BCrypt, filtres d'authentification, gestion des roles, suivi des "
                     "tentatives de connexion."],
        ["Temps reel", "WebSocket (STOMP) pour les notifications instantanees."],
        ["Base de donnees", "MySQL 8 via JPA / Hibernate."],
        ["Reporting", "JasperReports (modeles JRXML)."],
        ["Frontend", "Angular 21 (composants standalone), TypeScript 5.9, RxJS, Chart.js, layouts par role."],
        ["Documentation", "Swagger / OpenAPI 3.0."],
        ["Outils & build", "Maven (backend), Angular CLI / npm (frontend), Git."],
        ["Integrations", "Service d'email (SMTP), configuration Google Calendar."],
    ],
    col_widths=[1.7, 4.8],
)
h2("Principes d'architecture")
bullet("Architecture decouplee front / back communiquant par API REST (JSON).", )
bullet("Separation stricte des responsabilites en couches cote serveur.")
bullet("Interface monopage (SPA) avec routage et garde d'acces (guards) par role.")
bullet("Intercepteurs HTTP pour l'injection du jeton et la gestion centralisee des erreurs.")

# ===========================================================================
# 7. AXES D'INNOVATION
# ===========================================================================
h1("Axes d'innovation a developper", "7")
para("Au-dela du socle existant, les axes suivants visent a transformer la plateforme en un outil "
     "intelligent et differenciant. Ils constituent la valeur ajoutee innovante du projet.",
     italic=True, color=GREY)

innovations = [
    ("Assistant IA & priorisation intelligente",
     "Integration d'un modele d'IA (ex. API Claude) pour suggerer l'affectation optimale des taches, "
     "estimer la duree, prioriser automatiquement le backlog et resumer l'avancement d'un projet en "
     "langage naturel.", "Eleve"),
    ("Detection des risques & prediction de retard",
     "Analyse predictive de la sante des projets : alerte anticipee sur les depassements d'echeance et "
     "les surcharges d'equipe a partir de l'historique des taches et du suivi du temps.", "Eleve"),
    ("Tableau de bord analytique avance",
     "Indicateurs de performance (velocite, taux de respect des delais, burndown, charge par membre) et "
     "visualisations interactives enrichies.", "Moyen"),
    ("Recherche globale & filtres intelligents",
     "Barre de recherche unifiee (projets, taches, fichiers, messages) avec filtres avances et raccourcis.",
     "Moyen"),
    ("Application mobile / PWA",
     "Version Progressive Web App installable et hors-ligne, ou application mobile dediee, avec "
     "notifications push.", "Moyen"),
    ("Automatisations & workflows",
     "Regles configurables (ex. : a l'approbation d'un livrable, cloturer la tache et notifier le PM) et "
     "rappels automatiques d'echeance.", "Moyen"),
    ("Collaboration documentaire en temps reel",
     "Edition / annotation partagee des livrables et indicateurs de presence en ligne.", "Faible"),
    ("Internationalisation (i18n)",
     "Support multilingue (francais / anglais) de l'interface.", "Faible"),
]
make_table(
    ["Innovation", "Description", "Priorite"],
    [[i[0], i[1], i[2]] for i in innovations],
    col_widths=[1.7, 4.0, 0.8],
)

# ===========================================================================
# 8. A AJOUTER (BACKLOG)
# ===========================================================================
h1("Fonctionnalites a ajouter (backlog)", "8")
para("Les elements suivants completent le produit pour atteindre un niveau de maturite professionnel "
     "et industriel.")

h2("Fonctionnel")
for t in [
    "Diagramme de Gantt et chronologie des projets (timeline) avec dependances entre taches.",
    "Messagerie de groupe par projet (en complement des discussions individuelles existantes).",
    "Sous-taches et listes de controle (checklists) au sein d'une tache.",
    "Modeles de projet et de taches reutilisables pour acceleration du demarrage.",
    "Etiquettes (tags) et categorisation transversale des taches et projets.",
    "Synchronisation bidirectionnelle effective avec Google Calendar / Outlook.",
    "Gestion documentaire centralisee avec versionnage des fichiers.",
    "Module de feuilles de temps approuvables et facturation indicative.",
]:
    bullet(t)

h2("Technique & qualite")
for t in [
    "Couverture de tests automatises (unitaires, integration, end-to-end).",
    "Pipeline CI/CD (build, tests, deploiement automatise).",
    "Conteneurisation Docker et orchestration pour le deploiement.",
    "Pagination, mise en cache et optimisation des requetes a grande echelle.",
    "Journalisation centralisee et supervision (monitoring, alerting).",
    "Politique de sauvegarde et plan de reprise d'activite (PRA).",
    "Renforcement securite : double authentification (2FA), limitation de debit (rate limiting), audit RGPD.",
]:
    bullet(t)

# ===========================================================================
# 9. EXIGENCES NON FONCTIONNELLES
# ===========================================================================
h1("Exigences non fonctionnelles", "9")
make_table(
    ["Categorie", "Exigence"],
    [
        ["Performance", "Temps de reponse des pages < 2 s ; reponses API < 500 ms en conditions nominales."],
        ["Securite", "Chiffrement des mots de passe (BCrypt), jetons JWT, controle d'acces par role, "
                      "protection CSRF/XSS/injection SQL."],
        ["Disponibilite", "Objectif de disponibilite >= 99 % ; sauvegardes regulieres."],
        ["Scalabilite", "Architecture decouplee permettant la montee en charge horizontale."],
        ["Ergonomie", "Interface responsive, coherente (charte 3 couleurs), accessible et intuitive."],
        ["Maintenabilite", "Code modulaire en couches, documentation API, conventions de nommage."],
        ["Compatibilite", "Navigateurs modernes (Chrome, Firefox, Edge, Safari)."],
        ["Tracabilite", "Journalisation des actions sensibles et des connexions."],
    ],
    col_widths=[1.7, 4.8],
)

# ===========================================================================
# 10. PLANNING
# ===========================================================================
h1("Planning previsionnel et lots de livraison", "10")
make_table(
    ["Lot", "Contenu", "Statut"],
    [
        ["Lot 0 — Socle", "Authentification, utilisateurs, projets, taches, equipes, livrables, "
                          "calendrier, messagerie, notifications, rapports, support.", "Realise"],
        ["Lot 1 — Analytique", "Tableau de bord analytique avance, recherche globale, exports enrichis.",
         "A planifier"],
        ["Lot 2 — Intelligence", "Assistant IA, priorisation intelligente, detection des risques et "
                                 "prediction de retard.", "A planifier"],
        ["Lot 3 — Collaboration+", "Gantt, messagerie de groupe, sous-taches, automatisations, "
                                   "synchronisation calendrier.", "A planifier"],
        ["Lot 4 — Industrialisation", "Tests, CI/CD, Docker, supervision, 2FA, PWA/mobile.",
         "A planifier"],
    ],
    col_widths=[1.7, 4.0, 0.8],
)

# ===========================================================================
# 11. RISQUES
# ===========================================================================
h1("Risques et mesures de mitigation", "11")
make_table(
    ["Risque", "Impact", "Mesure de mitigation"],
    [
        ["Derive du perimetre", "Eleve", "Lotissement clair, priorisation et validation par lot."],
        ["Dette technique", "Moyen", "Tests automatises, revues de code, refactoring continu."],
        ["Securite des donnees", "Eleve", "Audits, 2FA, chiffrement, conformite RGPD."],
        ["Adoption utilisateur", "Moyen", "UX soignee, formation, documentation et accompagnement."],
        ["Cout d'integration IA", "Moyen", "Cadrage des cas d'usage, prototypage, suivi des couts d'API."],
    ],
    col_widths=[1.8, 0.9, 3.8],
)

# ===========================================================================
# 12. CONCLUSION
# ===========================================================================
h1("Criteres de reception et conclusion", "12")
h2("Criteres de reception")
for t in [
    "Conformite fonctionnelle de chaque lot au present cahier des charges.",
    "Absence d'anomalie bloquante en recette ; anomalies majeures corrigees.",
    "Respect des exigences non fonctionnelles (performance, securite).",
    "Documentation technique et utilisateur a jour.",
]:
    bullet(t)

h2("Conclusion")
para("La plateforme GPI App dispose aujourd'hui d'un socle fonctionnel complet et operationnel couvrant "
     "l'ensemble du cycle de gestion de projet, soutenu par une architecture technique moderne et "
     "evolutive. Les axes d'innovation identifies — assistant IA, analyse predictive des risques, "
     "tableau de bord analytique avance et industrialisation — constituent la feuille de route pour "
     "transformer ce socle en une solution intelligente, differenciante et prete pour un usage a grande "
     "echelle. Le lotissement propose permet une mise en oeuvre progressive, maitrisee et alignee sur les "
     "priorites de l'organisation.")

# Bandeau de fin
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
add_bottom_border(p, NAVY_HEX, 24)
para("Fin du document — Cahier des charges GPI App, version 1.0 — 29 mai 2026.",
     italic=True, color=GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------------------------------------------------------------------------
# Pied de page avec numerotation
# ---------------------------------------------------------------------------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Cahier des Charges — GPI App  |  Page ")
run.font.size = Pt(8)
run.font.color.rgb = GREY
run.font.name = FONT
# champ PAGE
fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
fp._p.append(fld1)

out = r"c:\Users\DON JUNIOR\Documents\Project\taskManagement-\docs\Cahier_des_Charges_GPI_App.docx"
doc.save(out)
print("OK ->", out)
